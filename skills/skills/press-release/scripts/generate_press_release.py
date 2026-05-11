#!/usr/bin/env python3
"""
generate_press_release.py

Generates a newswire-ready .docx press release from a structured JSON payload.

Usage:
    python3 generate_press_release.py --input payload.json --output release.docx

The payload JSON shape is:

{
  "release_type": "product_launch | funding | partnership | customer_win | exec_hire | milestone | award | acquisition",
  "embargo": null,                              # null = FOR IMMEDIATE RELEASE
                                                # or {"until": "Tuesday, May 12, 2026, 9:00 a.m. ET"}
  "headline": "social.plus Launches Commerce: Turning Community Into the New Storefront",
  "subhead":  "New capabilities bring product discovery directly into community posts and livestreams, closing the gap between high-intent conversation and action",
  "dateline": {"city": "LONDON", "date": "April 29, 2026"},
  "lede": "social.plus, the community infrastructure platform powering in-app social experiences for some of the world's leading brands, announced the launch of Commerce, a suite of capabilities that connects high-intent community moments to the products, services, and offerings users can actually act on, without leaving the app.",
  "second_para": "Commerce is now available across social.plus and includes three interconnected capabilities: Product Catalogue, Product Tagging in Posts, and Tag and Pin Products in Livestream.",

  "sections": [
    {
      "subhead": "The gap Commerce is built to close",
      "paragraphs": [
        "Community platforms generate some of the highest-intent signals in digital commerce. ...",
        "The result: engagement happens inside the community, but conversion happens somewhere else. ...",
        "Commerce changes this. ..."
      ]
    },
    {
      "subhead": null,                          # null subhead = just the quote block
      "quote": {
        "text": "Until commerce is a core part of the community experience, ...",
        "speaker_name": "Amadeus Norén",
        "speaker_title": "VP of Marketing"
      }
    },
    {
      "subhead": "What's new",
      "paragraphs": [
        {"lead": "Product Catalogue", "body": "is the foundation of Commerce. Brands can upload ..."},
        {"lead": "Product Tagging in Posts", "body": "allows brands and moderators to tag ..."},
        {"lead": "Tag and Pin Products in Livestream", "body": "brings shoppable experiences ..."}
      ]
    },
    {
      "subhead": "Built for every industry, not just retail",
      "paragraphs": [
        "While Commerce is purpose-built for retail and e-commerce brands ...",
        "Commerce is designed to close the distance between intent and action ..."
      ]
    },
    {
      "subhead": "Availability",
      "paragraphs": [
        "social.plus Commerce is available now across the social.plus platform. ..."
      ]
    }
  ],

  "boilerplate": {
    "title": "About social.plus",
    "body": "social.plus provides the community infrastructure ..."
  },
  "media_contact": "marketing@social.plus | social.plus",

  # Optional second boilerplate (for partnerships / acquisitions)
  "second_boilerplate": null
}

Each `sections[].paragraphs[]` item may be:
  - a plain string (regular paragraph)
  - {"lead": "X", "body": "..."} (lead phrase rendered bold, body follows)
  - {"quote": {...}} inside a section (rare; quotes usually live in their own section)
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches


# ----- Style constants -----------------------------------------------------

FONT_FAMILY = "Calibri"          # universally available; survives the wire's conversions
BODY_SIZE   = Pt(11)
HEADLINE_SIZE = Pt(18)
SUBHEAD_SIZE  = Pt(12)
SECTION_SUBHEAD_SIZE = Pt(12)
QUOTE_SIZE = Pt(11)

GREY_SUBHEAD = RGBColor(0x55, 0x55, 0x55)
BLACK        = RGBColor(0x00, 0x00, 0x00)


# ----- Helpers -------------------------------------------------------------

def _set_font(run, *, name=FONT_FAMILY, size=BODY_SIZE, bold=False, italic=False,
              color=BLACK):
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def _add_paragraph(doc, text="", *, bold=False, italic=False, size=BODY_SIZE,
                   color=BLACK, alignment=None, space_after=Pt(8)):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    if text:
        run = p.add_run(text)
        _set_font(run, size=size, bold=bold, italic=italic, color=color)
    p.paragraph_format.space_after = space_after
    return p


def _add_release_header(doc, embargo):
    """FOR IMMEDIATE RELEASE or UNDER EMBARGO block at the very top."""
    if embargo and embargo.get("until"):
        text = f"UNDER EMBARGO UNTIL {embargo['until']}"
    else:
        text = "FOR IMMEDIATE RELEASE"
    _add_paragraph(doc, text, bold=True, size=Pt(11),
                   alignment=WD_ALIGN_PARAGRAPH.LEFT,
                   space_after=Pt(18))


def _add_headline(doc, headline):
    _add_paragraph(doc, headline, bold=True, size=HEADLINE_SIZE,
                   alignment=WD_ALIGN_PARAGRAPH.LEFT,
                   space_after=Pt(6))


def _add_subhead(doc, subhead):
    if not subhead:
        return
    _add_paragraph(doc, subhead, italic=True, size=SUBHEAD_SIZE,
                   color=GREY_SUBHEAD,
                   alignment=WD_ALIGN_PARAGRAPH.LEFT,
                   space_after=Pt(18))


def _add_dateline_lede(doc, dateline, lede):
    """Combined paragraph: 'CITY, Month DD, YYYY — lede text...'"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)

    city = dateline.get("city", "").upper()
    date = dateline.get("date", "")
    dateline_text = f"{city}, {date} — "

    run_dl = p.add_run(dateline_text)
    _set_font(run_dl, bold=True)

    run_lede = p.add_run(lede)
    _set_font(run_lede, bold=False)


def _add_second_para(doc, text):
    if not text:
        return
    _add_paragraph(doc, text, space_after=Pt(14))


def _add_section_subhead(doc, text):
    if not text:
        return
    _add_paragraph(doc, text, bold=True, size=SECTION_SUBHEAD_SIZE,
                   alignment=WD_ALIGN_PARAGRAPH.LEFT,
                   space_after=Pt(8))


def _add_body_paragraph(doc, item):
    """item may be a string OR {'lead': '...', 'body': '...'}"""
    if isinstance(item, str):
        _add_paragraph(doc, item, space_after=Pt(10))
        return
    if isinstance(item, dict) and "lead" in item and "body" in item:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(10)
        run_lead = p.add_run(item["lead"])
        _set_font(run_lead, bold=True)
        # Single leading space between the lead phrase and the body
        body_text = item["body"]
        if not body_text.startswith(" "):
            body_text = " " + body_text
        run_body = p.add_run(body_text)
        _set_font(run_body, bold=False)
        return
    # Unknown shape — skip silently rather than corrupt the doc
    return


def _add_quote_block(doc, quote):
    """Indented italic quote paragraph followed by attribution on its own line."""
    if not quote or not quote.get("text"):
        return

    # Quote paragraph — indented, with curly quotes
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)

    quote_text = quote["text"]
    # Wrap with curly quotes if not already
    if not quote_text.startswith("“"):
        quote_text = "“" + quote_text
    if not quote_text.endswith("”"):
        quote_text = quote_text + "”"

    run = p.add_run(quote_text)
    _set_font(run, italic=False)

    # Attribution line — italic, indented, smaller
    speaker = quote.get("speaker_name", "")
    title = quote.get("speaker_title", "")
    affiliation = quote.get("speaker_affiliation")  # for customer/partner quotes
    attribution_parts = [x for x in [speaker, title, affiliation] if x]
    attribution = ", ".join(attribution_parts)

    pa = doc.add_paragraph()
    pa.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pa.paragraph_format.left_indent = Inches(0.4)
    pa.paragraph_format.right_indent = Inches(0.4)
    pa.paragraph_format.space_before = Pt(0)
    pa.paragraph_format.space_after = Pt(14)

    run_att = pa.add_run(attribution)
    _set_font(run_att, italic=True)


def _add_section(doc, section):
    if section.get("subhead"):
        _add_section_subhead(doc, section["subhead"])

    # Paragraphs
    for para in section.get("paragraphs", []) or []:
        _add_body_paragraph(doc, para)

    # Direct quote in section (some sections are just a quote block)
    if "quote" in section and section["quote"]:
        _add_quote_block(doc, section["quote"])


def _add_boilerplate(doc, boilerplate):
    if not boilerplate:
        return
    title = boilerplate.get("title", "About social.plus")
    body = boilerplate.get("body", "")
    _add_section_subhead(doc, title)
    _add_paragraph(doc, body, space_after=Pt(12))


def _add_media_contact(doc, contact):
    if not contact:
        return
    _add_section_subhead(doc, "Media Contact")
    _add_paragraph(doc, contact, space_after=Pt(20))


def _add_end_marker(doc):
    _add_paragraph(doc, "###", bold=True, size=Pt(12),
                   alignment=WD_ALIGN_PARAGRAPH.CENTER,
                   space_after=Pt(0))


# ----- Page setup ----------------------------------------------------------

def _setup_document_defaults(doc):
    """Set base font and conservative margins."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_FAMILY
    font.size = BODY_SIZE
    # No font.color reset — defaults to automatic black

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


# ----- Main ----------------------------------------------------------------

def build_press_release(payload, output_path):
    doc = Document()
    _setup_document_defaults(doc)

    _add_release_header(doc, payload.get("embargo"))
    _add_headline(doc, payload.get("headline", ""))
    _add_subhead(doc, payload.get("subhead"))
    _add_dateline_lede(doc, payload.get("dateline", {}), payload.get("lede", ""))
    _add_second_para(doc, payload.get("second_para"))

    for section in payload.get("sections", []) or []:
        _add_section(doc, section)

    _add_boilerplate(doc, payload.get("boilerplate"))
    _add_boilerplate(doc, payload.get("second_boilerplate"))
    _add_media_contact(doc, payload.get("media_contact"))
    _add_end_marker(doc)

    doc.save(output_path)


def main():
    ap = argparse.ArgumentParser(description="Generate a newswire-ready press release .docx")
    ap.add_argument("--input", "-i", required=True, help="Path to JSON payload")
    ap.add_argument("--output", "-o", required=True, help="Path to write the .docx")
    args = ap.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)

    if not input_path.is_file():
        print(f"ERROR: input file not found: {input_path}", file=sys.stderr)
        sys.exit(2)

    with input_path.open("r", encoding="utf-8") as f:
        payload = json.load(f)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    build_press_release(payload, str(output_path))
    print(f"Wrote: {output_path}")


if __name__ == "__main__":
    main()
